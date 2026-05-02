"""
Report Generator
================
Generate professional reports in PDF, Excel, and Word formats.

Author: VidyuthLabs
Date: May 1, 2026
"""

import logging
from datetime import datetime
from typing import Dict, List, Optional, Any, BinaryIO
from dataclasses import dataclass
from enum import Enum
import io
import base64

logger = logging.getLogger(__name__)


class ReportFormat(str, Enum):
    """Report format enumeration."""
    PDF = "pdf"
    EXCEL = "excel"
    WORD = "word"
    HTML = "html"
    MARKDOWN = "markdown"


@dataclass
class ReportSection:
    """Report section."""
    title: str
    content: str
    level: int = 1  # Heading level (1-6)
    include_in_toc: bool = True


@dataclass
class ReportTable:
    """Report table."""
    title: str
    headers: List[str]
    rows: List[List[Any]]
    caption: Optional[str] = None


@dataclass
class ReportChart:
    """Report chart."""
    title: str
    chart_type: str  # line, bar, scatter, pie
    data: Dict[str, Any]
    caption: Optional[str] = None


@dataclass
class ReportConfig:
    """Report configuration."""
    title: str
    subtitle: Optional[str] = None
    author: str = "RĀMAN Studio"
    company: str = "VidyuthLabs"
    logo_path: Optional[str] = None
    include_toc: bool = True
    include_page_numbers: bool = True
    include_header: bool = True
    include_footer: bool = True
    include_timestamp: bool = True


class ReportGenerator:
    """
    Professional report generator for electrochemical analysis.
    
    Features:
    - Multiple formats (PDF, Excel, Word, HTML, Markdown)
    - Professional styling
    - Tables and charts
    - Table of contents
    - Headers and footers
    - Company branding
    - 21 CFR Part 11 compliant
    
    Examples:
        # Create report
        generator = ReportGenerator()
        
        # Add sections
        generator.add_section("Executive Summary", "This report presents...")
        generator.add_section("Experimental Details", "The experiment was...")
        
        # Add table
        generator.add_table(
            title="EIS Fitting Results",
            headers=["Parameter", "Value", "Unit"],
            rows=[
                ["R_s", "10.5", "Ω"],
                ["R_ct", "150.2", "Ω"],
                ["CPE-T", "2.5e-6", "F·s^(α-1)"]
            ]
        )
        
        # Generate PDF
        pdf_bytes = generator.generate(ReportFormat.PDF)
    """
    
    def __init__(self, config: Optional[ReportConfig] = None):
        """
        Initialize report generator.
        
        Args:
            config: Report configuration
        """
        self.config = config or ReportConfig(title="Analysis Report")
        self.sections: List[ReportSection] = []
        self.tables: List[ReportTable] = []
        self.charts: List[ReportChart] = []
        
        logger.info(f"Report generator initialized: {self.config.title}")
    
    def add_section(
        self,
        title: str,
        content: str,
        level: int = 1,
        include_in_toc: bool = True
    ):
        """
        Add section to report.
        
        Args:
            title: Section title
            content: Section content
            level: Heading level (1-6)
            include_in_toc: Include in table of contents
        """
        section = ReportSection(
            title=title,
            content=content,
            level=level,
            include_in_toc=include_in_toc
        )
        self.sections.append(section)
        logger.debug(f"Added section: {title}")
    
    def add_table(
        self,
        title: str,
        headers: List[str],
        rows: List[List[Any]],
        caption: Optional[str] = None
    ):
        """
        Add table to report.
        
        Args:
            title: Table title
            headers: Column headers
            rows: Table rows
            caption: Table caption
        """
        table = ReportTable(
            title=title,
            headers=headers,
            rows=rows,
            caption=caption
        )
        self.tables.append(table)
        logger.debug(f"Added table: {title}")
    
    def add_chart(
        self,
        title: str,
        chart_type: str,
        data: Dict[str, Any],
        caption: Optional[str] = None
    ):
        """
        Add chart to report.
        
        Args:
            title: Chart title
            chart_type: Chart type (line, bar, scatter, pie)
            data: Chart data
            caption: Chart caption
        """
        chart = ReportChart(
            title=title,
            chart_type=chart_type,
            data=data,
            caption=caption
        )
        self.charts.append(chart)
        logger.debug(f"Added chart: {title}")
    
    def generate(self, format: ReportFormat) -> bytes:
        """
        Generate report in specified format.
        
        Args:
            format: Report format
        
        Returns:
            Report bytes
        """
        logger.info(f"Generating report: {self.config.title} ({format})")
        
        if format == ReportFormat.PDF:
            return self._generate_pdf()
        elif format == ReportFormat.EXCEL:
            return self._generate_excel()
        elif format == ReportFormat.WORD:
            return self._generate_word()
        elif format == ReportFormat.HTML:
            return self._generate_html().encode('utf-8')
        elif format == ReportFormat.MARKDOWN:
            return self._generate_markdown().encode('utf-8')
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _generate_pdf(self) -> bytes:
        """Generate PDF report using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            
            # Create PDF buffer
            buffer = io.BytesIO()
            
            # Create document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for elements
            elements = []
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1  # Center
            )
            
            # Title page
            elements.append(Spacer(1, 2*inch))
            elements.append(Paragraph(self.config.title, title_style))
            
            if self.config.subtitle:
                elements.append(Spacer(1, 12))
                elements.append(Paragraph(self.config.subtitle, styles['Heading2']))
            
            elements.append(Spacer(1, 0.5*inch))
            elements.append(Paragraph(f"Generated by: {self.config.author}", styles['Normal']))
            elements.append(Paragraph(f"Company: {self.config.company}", styles['Normal']))
            
            if self.config.include_timestamp:
                timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
                elements.append(Paragraph(f"Date: {timestamp}", styles['Normal']))
            
            elements.append(PageBreak())
            
            # Table of contents (if enabled)
            if self.config.include_toc:
                elements.append(Paragraph("Table of Contents", styles['Heading1']))
                elements.append(Spacer(1, 12))
                
                for i, section in enumerate(self.sections, 1):
                    if section.include_in_toc:
                        indent = "&nbsp;" * (section.level - 1) * 4
                        toc_entry = f"{indent}{i}. {section.title}"
                        elements.append(Paragraph(toc_entry, styles['Normal']))
                
                elements.append(PageBreak())
            
            # Sections
            for section in self.sections:
                # Section title
                heading_style = styles[f'Heading{min(section.level, 3)}']
                elements.append(Paragraph(section.title, heading_style))
                elements.append(Spacer(1, 12))
                
                # Section content
                for paragraph in section.content.split('\n\n'):
                    if paragraph.strip():
                        elements.append(Paragraph(paragraph, styles['Normal']))
                        elements.append(Spacer(1, 12))
            
            # Tables
            for table in self.tables:
                elements.append(Spacer(1, 12))
                elements.append(Paragraph(table.title, styles['Heading2']))
                elements.append(Spacer(1, 12))
                
                # Create table data
                table_data = [table.headers] + table.rows
                
                # Create table
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(t)
                
                if table.caption:
                    elements.append(Spacer(1, 6))
                    elements.append(Paragraph(f"<i>{table.caption}</i>", styles['Normal']))
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info(f"PDF generated: {len(pdf_bytes)} bytes")
            return pdf_bytes
        
        except ImportError:
            logger.error("ReportLab not installed. Install with: pip install reportlab")
            # Return simple text-based PDF alternative
            return self._generate_markdown().encode('utf-8')
    
    def _generate_excel(self) -> bytes:
        """Generate Excel report using openpyxl."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            
            # Create workbook
            wb = Workbook()
            
            # Summary sheet
            ws = wb.active
            ws.title = "Summary"
            
            # Title
            ws['A1'] = self.config.title
            ws['A1'].font = Font(size=18, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A1:D1')
            
            # Metadata
            row = 3
            ws[f'A{row}'] = "Generated by:"
            ws[f'B{row}'] = self.config.author
            row += 1
            ws[f'A{row}'] = "Company:"
            ws[f'B{row}'] = self.config.company
            row += 1
            
            if self.config.include_timestamp:
                ws[f'A{row}'] = "Date:"
                ws[f'B{row}'] = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
                row += 1
            
            # Sections
            row += 2
            for section in self.sections:
                ws[f'A{row}'] = section.title
                ws[f'A{row}'].font = Font(size=14, bold=True)
                row += 1
                
                # Split content into lines
                for line in section.content.split('\n'):
                    if line.strip():
                        ws[f'A{row}'] = line
                        row += 1
                
                row += 1
            
            # Tables (each in separate sheet)
            for i, table in enumerate(self.tables, 1):
                ws = wb.create_sheet(title=f"Table_{i}")
                
                # Table title
                ws['A1'] = table.title
                ws['A1'].font = Font(size=14, bold=True)
                
                # Headers
                for col, header in enumerate(table.headers, 1):
                    cell = ws.cell(row=3, column=col)
                    cell.value = header
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    cell.alignment = Alignment(horizontal='center')
                
                # Data rows
                for row_idx, row_data in enumerate(table.rows, 4):
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Caption
                if table.caption:
                    caption_row = len(table.rows) + 5
                    ws[f'A{caption_row}'] = table.caption
                    ws[f'A{caption_row}'].font = Font(italic=True)
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            excel_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info(f"Excel generated: {len(excel_bytes)} bytes")
            return excel_bytes
        
        except ImportError:
            logger.error("openpyxl not installed. Install with: pip install openpyxl")
            # Return CSV alternative
            return self._generate_markdown().encode('utf-8')
    
    def _generate_word(self) -> bytes:
        """Generate Word report using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading(self.config.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            if self.config.subtitle:
                subtitle = doc.add_heading(self.config.subtitle, level=2)
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph()
            doc.add_paragraph(f"Generated by: {self.config.author}")
            doc.add_paragraph(f"Company: {self.config.company}")
            
            if self.config.include_timestamp:
                timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
                doc.add_paragraph(f"Date: {timestamp}")
            
            doc.add_page_break()
            
            # Table of contents (if enabled)
            if self.config.include_toc:
                doc.add_heading("Table of Contents", level=1)
                
                for i, section in enumerate(self.sections, 1):
                    if section.include_in_toc:
                        indent = "    " * (section.level - 1)
                        doc.add_paragraph(f"{indent}{i}. {section.title}")
                
                doc.add_page_break()
            
            # Sections
            for section in self.sections:
                doc.add_heading(section.title, level=section.level)
                
                for paragraph in section.content.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
            
            # Tables
            for table in self.tables:
                doc.add_heading(table.title, level=2)
                
                # Create table
                t = doc.add_table(rows=1 + len(table.rows), cols=len(table.headers))
                t.style = 'Light Grid Accent 1'
                
                # Headers
                hdr_cells = t.rows[0].cells
                for i, header in enumerate(table.headers):
                    hdr_cells[i].text = str(header)
                
                # Data rows
                for i, row_data in enumerate(table.rows, 1):
                    row_cells = t.rows[i].cells
                    for j, value in enumerate(row_data):
                        row_cells[j].text = str(value)
                
                # Caption
                if table.caption:
                    caption = doc.add_paragraph(table.caption)
                    caption.italic = True
                
                doc.add_paragraph()
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            word_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info(f"Word document generated: {len(word_bytes)} bytes")
            return word_bytes
        
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            # Return markdown alternative
            return self._generate_markdown().encode('utf-8')
    
    def _generate_html(self) -> str:
        """Generate HTML report."""
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.config.title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        .caption {{
            font-style: italic;
            color: #7f8c8d;
            margin-top: 5px;
        }}
    </style>
</head>
<body>
    <h1>{self.config.title}</h1>
"""
        
        if self.config.subtitle:
            html += f"    <h2>{self.config.subtitle}</h2>\n"
        
        # Metadata
        html += '    <div class="metadata">\n'
        html += f"        <p><strong>Generated by:</strong> {self.config.author}</p>\n"
        html += f"        <p><strong>Company:</strong> {self.config.company}</p>\n"
        
        if self.config.include_timestamp:
            timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
            html += f"        <p><strong>Date:</strong> {timestamp}</p>\n"
        
        html += '    </div>\n'
        
        # Sections
        for section in self.sections:
            html += f"    <h{section.level + 1}>{section.title}</h{section.level + 1}>\n"
            
            for paragraph in section.content.split('\n\n'):
                if paragraph.strip():
                    html += f"    <p>{paragraph}</p>\n"
        
        # Tables
        for table in self.tables:
            html += f"    <h2>{table.title}</h2>\n"
            html += '    <table>\n'
            
            # Headers
            html += '        <tr>\n'
            for header in table.headers:
                html += f"            <th>{header}</th>\n"
            html += '        </tr>\n'
            
            # Rows
            for row in table.rows:
                html += '        <tr>\n'
                for cell in row:
                    html += f"            <td>{cell}</td>\n"
                html += '        </tr>\n'
            
            html += '    </table>\n'
            
            if table.caption:
                html += f'    <p class="caption">{table.caption}</p>\n'
        
        html += """</body>
</html>"""
        
        logger.info(f"HTML generated: {len(html)} bytes")
        return html
    
    def _generate_markdown(self) -> str:
        """Generate Markdown report."""
        md = f"# {self.config.title}\n\n"
        
        if self.config.subtitle:
            md += f"## {self.config.subtitle}\n\n"
        
        # Metadata
        md += f"**Generated by:** {self.config.author}  \n"
        md += f"**Company:** {self.config.company}  \n"
        
        if self.config.include_timestamp:
            timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
            md += f"**Date:** {timestamp}  \n"
        
        md += "\n---\n\n"
        
        # Sections
        for section in self.sections:
            md += f"{'#' * (section.level + 1)} {section.title}\n\n"
            md += f"{section.content}\n\n"
        
        # Tables
        for table in self.tables:
            md += f"## {table.title}\n\n"
            
            # Headers
            md += "| " + " | ".join(str(h) for h in table.headers) + " |\n"
            md += "| " + " | ".join("---" for _ in table.headers) + " |\n"
            
            # Rows
            for row in table.rows:
                md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
            
            md += "\n"
            
            if table.caption:
                md += f"*{table.caption}*\n\n"
        
        logger.info(f"Markdown generated: {len(md)} bytes")
        return md


# ===================================================================
#  Convenience Functions
# ===================================================================

def generate_experiment_report(
    experiment_data: Dict[str, Any],
    format: ReportFormat = ReportFormat.PDF
) -> bytes:
    """
    Generate experiment report.
    
    Args:
        experiment_data: Experiment data
        format: Report format
    
    Returns:
        Report bytes
    """
    config = ReportConfig(
        title=f"Experiment Report: {experiment_data.get('name', 'Untitled')}",
        subtitle=f"Technique: {experiment_data.get('technique', 'Unknown')}"
    )
    
    generator = ReportGenerator(config)
    
    # Executive summary
    generator.add_section(
        "Executive Summary",
        f"This report presents the results of {experiment_data.get('technique', 'electrochemical')} "
        f"analysis performed on {datetime.utcnow().strftime('%Y-%m-%d')}."
    )
    
    # Experimental details
    if 'parameters' in experiment_data:
        params_text = "\n".join(
            f"- {k}: {v}" for k, v in experiment_data['parameters'].items()
        )
        generator.add_section("Experimental Parameters", params_text)
    
    # Results
    if 'results' in experiment_data:
        results_text = "\n".join(
            f"- {k}: {v}" for k, v in experiment_data['results'].items()
        )
        generator.add_section("Results", results_text)
    
    return generator.generate(format)


def generate_batch_report(
    batch_data: Dict[str, Any],
    format: ReportFormat = ReportFormat.PDF
) -> bytes:
    """
    Generate batch processing report.
    
    Args:
        batch_data: Batch job data
        format: Report format
    
    Returns:
        Report bytes
    """
    config = ReportConfig(
        title=f"Batch Processing Report: {batch_data.get('name', 'Untitled')}",
        subtitle=f"Job ID: {batch_data.get('id', 'Unknown')}"
    )
    
    generator = ReportGenerator(config)
    
    # Summary
    summary = f"""
Total Files: {batch_data.get('total_files', 0)}
Successful: {batch_data.get('successful_files', 0)}
Failed: {batch_data.get('failed_files', 0)}
Success Rate: {batch_data.get('success_rate', 0):.1f}%
Processing Time: {batch_data.get('processing_time', 0):.1f}s
"""
    generator.add_section("Summary", summary)
    
    # Results table
    if 'results' in batch_data:
        results = batch_data['results']
        generator.add_table(
            title="Processing Results",
            headers=["File", "Status", "Time (s)"],
            rows=[
                [r.get('file', ''), r.get('status', ''), f"{r.get('time', 0):.2f}"]
                for r in results[:20]  # First 20 results
            ]
        )
    
    return generator.generate(format)

#!/usr/bin/env python3
"""
export_document.py — Export Markdown+LaTeX+{{fig:ID}} report to DOCX or PDF.

Usage:
    python export_document.py <config_json_path>

Config JSON:
{
  "job_id": "...",
  "job_dir": "/tmp/rvce_jobs/...",
  "output_path": "/tmp/...",
  "format": "docx" | "pdf",
  "content": "...",
  "settings": {
    "title": "...",
    "authors": "...",
    "institution": "...",
    "font_family": "Times New Roman",
    "font_size": 11,
    "margins": "normal",
    "line_spacing": 1.5
  }
}

Outputs JSON to stdout: {ok, output_path, error?}
"""
from __future__ import annotations

import json, sys, re, io, os, traceback
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image as PILImage

# ── Math rendering ─────────────────────────────────────────────────────────────

def render_math_png(expr: str, fontsize: float = 14, display: bool = False) -> bytes | None:
    """Render LaTeX math to PNG bytes using matplotlib mathtext."""
    try:
        expr = expr.strip()
        fs = fontsize * (1.3 if display else 1.0)
        fig, ax = plt.subplots(figsize=(8, 0.8))
        ax.set_axis_off()
        ax.text(0.5, 0.5, f"${expr}$", transform=ax.transAxes,
                fontsize=fs, ha="center", va="center", color="black")
        fig.patch.set_facecolor("white")
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    pad_inches=0.08, facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None

# ── Content parser ─────────────────────────────────────────────────────────────

FIG_RE    = re.compile(r"\{\{fig:([a-z0-9_]+)\}\}")
DISP_MATH = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
INLN_MATH = re.compile(r"\$(.+?)\$")
TABLE_ROW = re.compile(r"^\|(.+)\|$")
HR_RE     = re.compile(r"^---+$")
OL_RE     = re.compile(r"^\d+\.\s+(.+)$")
UL_RE     = re.compile(r"^[-*]\s+(.+)$")

def _strip_inline(text: str) -> str:
    """Remove markdown bold/italic markers for plain-text use."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*",     r"\1", text)
    text = re.sub(r"\*(.+?)\*",         r"\1", text)
    return text

class Block:
    pass

class HeadingBlock(Block):
    def __init__(self, level: int, text: str): self.level = level; self.text = text

class ParagraphBlock(Block):
    def __init__(self, text: str): self.text = text

class DisplayMathBlock(Block):
    def __init__(self, expr: str): self.expr = expr

class FigureBlock(Block):
    def __init__(self, plot_id: str, image_path: Path | None, caption: str):
        self.plot_id = plot_id; self.image_path = image_path; self.caption = caption

class TableBlock(Block):
    def __init__(self, headers: list[str], rows: list[list[str]]):
        self.headers = headers; self.rows = rows

class ListBlock(Block):
    def __init__(self, items: list[str], ordered: bool):
        self.items = items; self.ordered = ordered

class RuleBlock(Block):
    pass

def find_plot_image(job_dir: Path, plot_id: str) -> Path | None:
    """Search for a plot image file in the job output directory."""
    for ext in (".png", ".jpg", ".jpeg"):
        # walk all sample dirs and comparison dirs
        for candidate in job_dir.rglob(f"{plot_id}{ext}"):
            if candidate.is_file():
                return candidate
    return None

def parse_content(text: str, job_dir: Path | None, figure_captions: dict[str, str]) -> list[Block]:
    """Parse markdown+LaTeX+{{fig:}} content into a list of Block objects."""
    blocks: list[Block] = []

    # Replace display math placeholders to protect them during line-splitting
    display_maths: list[str] = []
    def _save_disp(m: re.Match) -> str:
        display_maths.append(m.group(1))
        return f"\x00DISPMATH{len(display_maths)-1}\x00"
    text = DISP_MATH.sub(_save_disp, text)

    lines = text.split("\n")
    i = 0
    para_lines: list[str] = []

    def flush_para():
        nonlocal para_lines
        chunk = " ".join(l.strip() for l in para_lines if l.strip())
        if chunk:
            blocks.append(ParagraphBlock(chunk))
        para_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line → flush paragraph
        if not stripped:
            flush_para()
            i += 1
            continue

        # Display math placeholder
        dm = re.match(r"^\x00DISPMATH(\d+)\x00$", stripped)
        if dm:
            flush_para()
            blocks.append(DisplayMathBlock(display_maths[int(dm.group(1))]))
            i += 1
            continue

        # Figure reference
        fm = FIG_RE.match(stripped)
        if fm:
            flush_para()
            pid = fm.group(1)
            img = find_plot_image(job_dir, pid) if job_dir else None
            cap = figure_captions.get(pid, pid.replace("_", " ").title())
            blocks.append(FigureBlock(pid, img, cap))
            i += 1
            continue

        # Headings
        hm = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if hm:
            flush_para()
            blocks.append(HeadingBlock(len(hm.group(1)), _strip_inline(hm.group(2))))
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(stripped):
            flush_para()
            blocks.append(RuleBlock())
            i += 1
            continue

        # Table
        if TABLE_ROW.match(stripped):
            flush_para()
            headers: list[str] = []
            rows: list[list[str]] = []
            while i < len(lines) and TABLE_ROW.match(lines[i].strip()):
                cols = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if all(re.match(r"^[-:]+$", c) for c in cols):
                    i += 1
                    continue
                if not headers:
                    headers = cols
                else:
                    rows.append(cols)
                i += 1
            blocks.append(TableBlock(headers, rows))
            continue

        # Ordered list
        if OL_RE.match(stripped):
            flush_para()
            items: list[str] = []
            while i < len(lines) and OL_RE.match(lines[i].strip()):
                items.append(OL_RE.match(lines[i].strip()).group(1))
                i += 1
            blocks.append(ListBlock(items, ordered=True))
            continue

        # Unordered list
        if UL_RE.match(stripped):
            flush_para()
            items_u: list[str] = []
            while i < len(lines) and UL_RE.match(lines[i].strip()):
                items_u.append(UL_RE.match(lines[i].strip()).group(1))
                i += 1
            blocks.append(ListBlock(items_u, ordered=False))
            continue

        # Regular paragraph line
        para_lines.append(stripped)
        i += 1

    flush_para()
    return blocks

# ── PDF export (reportlab) ─────────────────────────────────────────────────────

def export_pdf(blocks: list[Block], output_path: Path, settings: dict) -> None:
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Image, Table, TableStyle, HRFlowable, PageBreak)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    pt = 1  # reportlab's native unit is points
    from reportlab.lib.colors import black, HexColor, white
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    from reportlab.lib.pagesizes import A4

    font_size    = float(settings.get("font_size", 11))
    margins_name = settings.get("margins", "normal")
    margin_map   = {"narrow": 1.27*cm, "normal": 2.54*cm, "wide": 3.81*cm}
    margin       = margin_map.get(margins_name, 2.54*cm)
    line_sp      = float(settings.get("line_spacing", 1.5))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin,  bottomMargin=margin,
    )

    styles = getSampleStyleSheet()
    body_font = "Times-Roman"
    bold_font = "Times-Bold"
    ital_font = "Times-Italic"

    ST = {}
    ST["body"] = ParagraphStyle("body", fontName=body_font, fontSize=font_size,
        leading=font_size * line_sp, spaceBefore=4, spaceAfter=6,
        alignment=TA_JUSTIFY)
    ST["h1"] = ParagraphStyle("h1", fontName=bold_font, fontSize=font_size+5,
        leading=(font_size+5)*1.2, spaceBefore=18, spaceAfter=10, alignment=TA_CENTER)
    ST["h2"] = ParagraphStyle("h2", fontName=bold_font, fontSize=font_size+2,
        leading=(font_size+2)*1.2, spaceBefore=14, spaceAfter=6)
    ST["h3"] = ParagraphStyle("h3", fontName=ital_font, fontSize=font_size+1,
        leading=(font_size+1)*1.2, spaceBefore=10, spaceAfter=4)
    ST["h4"] = ParagraphStyle("h4", fontName=bold_font, fontSize=font_size,
        leading=font_size*1.2, spaceBefore=8, spaceAfter=3)
    ST["caption"] = ParagraphStyle("caption", fontName=ital_font, fontSize=font_size-2,
        leading=(font_size-2)*1.3, spaceBefore=4, spaceAfter=8, alignment=TA_CENTER)
    ST["list_item"] = ParagraphStyle("list_item", fontName=body_font, fontSize=font_size,
        leading=font_size*line_sp, leftIndent=18, spaceBefore=2, spaceAfter=2)
    ST["meta"] = ParagraphStyle("meta", fontName=ital_font, fontSize=font_size,
        leading=font_size*1.3, spaceBefore=4, spaceAfter=2, alignment=TA_CENTER)

    title     = settings.get("title", "Electrochemical CV Analysis Report")
    authors   = settings.get("authors", "")
    institute = settings.get("institution", "")

    def _safe_para(text: str, style: ParagraphStyle) -> Paragraph:
        """Escape special chars and handle inline math by rendering to images."""
        # Replace inline math $...$ with [MATH] placeholder (simplified for PDF)
        def repl_inline(m: re.Match) -> str:
            expr = m.group(1)
            # Use unicode approximations for common symbols
            expr = (expr.replace(r"\bar", "")
                        .replace(r"\nu", "ν").replace(r"\alpha", "α")
                        .replace(r"\beta", "β").replace(r"\gamma", "γ")
                        .replace(r"\delta", "δ").replace(r"\sigma", "σ")
                        .replace(r"\mu", "μ").replace(r"\pi", "π")
                        .replace(r"\omega", "ω").replace(r"\theta", "θ")
                        .replace(r"\infty", "∞").replace(r"\pm", "±")
                        .replace(r"\leq", "≤").replace(r"\geq", "≥")
                        .replace(r"\approx", "≈").replace(r"\sqrt", "√")
                        .replace(r"\times", "×").replace(r"\cdot", "·")
                        .replace(r"\text{", "").replace("}", "")
                        .replace("{", "").replace("}", ""))
            return f"<i>{expr}</i>"
        text = INLN_MATH.sub(repl_inline, text)
        # Escape XML special chars but allow our italic tags
        text = (text.replace("&", "&amp;").replace("<i>", "\x01i\x01")
                    .replace("</i>", "\x01/i\x01").replace("<", "&lt;").replace(">", "&gt;")
                    .replace("\x01i\x01", "<i>").replace("\x01/i\x01", "</i>"))
        # Markdown bold/italic
        text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
        text = re.sub(r"\*\*(.+?)\*\*",     r"<b>\1</b>",         text)
        text = re.sub(r"\*(.+?)\*",         r"<i>\1</i>",         text)
        try:
            return Paragraph(text, style)
        except Exception:
            return Paragraph(re.sub("<[^>]+>", "", text), style)

    story: list[Any] = []

    # ── Title block
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(title, ST["h1"]))
    if authors:
        story.append(Paragraph(authors, ST["meta"]))
    if institute:
        story.append(Paragraph(institute, ST["meta"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#888888")))
    story.append(Spacer(1, 0.4*cm))

    page_w = A4[0] - 2 * margin
    fig_num = [0]

    for block in blocks:
        if isinstance(block, HeadingBlock):
            key = f"h{block.level}"
            story.append(_safe_para(block.text, ST.get(key, ST["h2"])))

        elif isinstance(block, ParagraphBlock):
            story.append(_safe_para(block.text, ST["body"]))

        elif isinstance(block, DisplayMathBlock):
            png = render_math_png(block.expr, fontsize=font_size * 1.1, display=True)
            if png:
                img_obj = Image(io.BytesIO(png))
                w, h = img_obj.imageWidth, img_obj.imageHeight
                scale = min(page_w / w, 3*cm / h, 1.0)
                img_obj.drawWidth  = w * scale
                img_obj.drawHeight = h * scale
                story.append(Spacer(1, 0.2*cm))
                story.append(img_obj)
                story.append(Spacer(1, 0.2*cm))
            else:
                story.append(_safe_para(block.expr, ST["body"]))

        elif isinstance(block, FigureBlock):
            if block.image_path and block.image_path.exists():
                fig_num[0] += 1
                img_obj = Image(str(block.image_path))
                w, h = img_obj.imageWidth, img_obj.imageHeight
                max_w = page_w
                max_h = 14 * cm
                scale = min(max_w / w, max_h / h, 1.0)
                img_obj.drawWidth  = w * scale
                img_obj.drawHeight = h * scale
                story.append(Spacer(1, 0.4*cm))
                story.append(img_obj)
                cap = f"Figure {fig_num[0]}: {block.caption}"
                story.append(Paragraph(cap, ST["caption"]))
                story.append(Spacer(1, 0.3*cm))

        elif isinstance(block, TableBlock):
            if block.headers or block.rows:
                tdata = [block.headers] + block.rows if block.headers else block.rows
                col_w = page_w / max(len(tdata[0]), 1) if tdata else page_w
                tbl = Table(tdata, colWidths=[col_w]*len(tdata[0]))
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",   (0, 0), (-1, 0),  HexColor("#e8e8e8")),
                    ("FONTNAME",     (0, 0), (-1, 0),  bold_font),
                    ("FONTSIZE",     (0, 0), (-1, -1), font_size - 1),
                    ("FONTNAME",     (0, 1), (-1, -1), body_font),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f7f7f7")]),
                    ("GRID",         (0, 0), (-1, -1), 0.5, HexColor("#bbbbbb")),
                    ("ALIGN",        (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING",   (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
                    ("LEFTPADDING",  (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(Spacer(1, 0.2*cm))
                story.append(tbl)
                story.append(Spacer(1, 0.2*cm))

        elif isinstance(block, ListBlock):
            prefix_fn = (lambda i: f"{i+1}. ") if block.ordered else (lambda i: "• ")
            for idx, item in enumerate(block.items):
                story.append(_safe_para(f"{prefix_fn(idx)}{item}", ST["list_item"]))

        elif isinstance(block, RuleBlock):
            story.append(Spacer(1, 0.2*cm))
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#aaaaaa")))
            story.append(Spacer(1, 0.2*cm))

    doc.build(story)

# ── DOCX export (python-docx) ──────────────────────────────────────────────────

def export_docx(blocks: list[Block], output_path: Path, settings: dict) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    font_name  = settings.get("font_family", "Times New Roman")
    font_size  = float(settings.get("font_size", 11))
    margins_nm = settings.get("margins", "normal")
    margin_map = {"narrow": 1.27, "normal": 2.54, "wide": 3.81}
    margin_cm  = margin_map.get(margins_nm, 2.54)
    line_sp    = float(settings.get("line_spacing", 1.5))
    title_txt  = settings.get("title", "Electrochemical CV Analysis Report")
    authors    = settings.get("authors", "")
    institute  = settings.get("institution", "")

    doc = Document()

    # Page margins
    for section in doc.sections:
        m = Cm(margin_cm)
        section.left_margin = section.right_margin = m
        section.top_margin  = section.bottom_margin = m

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)

    def _set_font(run, bold=False, italic=False, size=None):
        run.font.name = font_name
        run.font.size = Pt(size or font_size)
        run.font.bold = bold
        run.font.italic = italic

    def _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
        fmt = para.paragraph_format
        fmt.alignment = alignment
        fmt.space_before = Pt(space_before)
        fmt.space_after  = Pt(space_after)
        fmt.line_spacing = Pt(font_size * line_sp)

    def _add_text_para(text: str, bold=False, italic=False, size=None,
                        align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6) -> None:
        """Add a paragraph with inline math/bold/italic parsing."""
        para = doc.add_paragraph()
        _set_para_format(para, alignment=align, space_before=before, space_after=after)
        # Split on inline math
        parts = INLN_MATH.split(text)
        is_math = False
        for part in parts:
            if is_math:
                # Render math as small inline image
                png = render_math_png(part, fontsize=font_size, display=False)
                if png:
                    run = para.add_run()
                    run.add_picture(io.BytesIO(png), height=Pt(font_size * 1.1))
                else:
                    # Fallback: plain text approximation
                    run = para.add_run(_unicode_math(part))
                    _set_font(run, italic=True, size=size)
            else:
                # Handle bold/italic markdown
                segments = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", part)
                for seg in segments:
                    if re.match(r"^\*\*\*.+\*\*\*$", seg):
                        run = para.add_run(seg[3:-3])
                        _set_font(run, bold=True, italic=True, size=size)
                    elif re.match(r"^\*\*.+\*\*$", seg):
                        run = para.add_run(seg[2:-2])
                        _set_font(run, bold=bold or True, italic=italic, size=size)
                    elif re.match(r"^\*.+\*$", seg):
                        run = para.add_run(seg[1:-1])
                        _set_font(run, bold=bold, italic=True, size=size)
                    else:
                        if seg:
                            run = para.add_run(seg)
                            _set_font(run, bold=bold, italic=italic, size=size)
            is_math = not is_math

    def _unicode_math(expr: str) -> str:
        return (expr.replace(r"\bar", "").replace(r"\nu", "ν").replace(r"\alpha", "α")
                    .replace(r"\beta", "β").replace(r"\gamma", "γ").replace(r"\delta", "δ")
                    .replace(r"\sigma", "σ").replace(r"\mu", "μ").replace(r"\pi", "π")
                    .replace(r"\omega", "ω").replace(r"\theta", "θ").replace(r"\infty", "∞")
                    .replace(r"\pm", "±").replace(r"\leq", "≤").replace(r"\geq", "≥")
                    .replace(r"\approx", "≈").replace(r"\sqrt", "√").replace(r"\times", "×")
                    .replace(r"\cdot", "·").replace(r"\text{", "").replace("{", "").replace("}", ""))

    # Title block
    para = doc.add_paragraph()
    _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
    run = para.add_run(title_txt)
    _set_font(run, bold=True, size=font_size + 4)

    if authors:
        para = doc.add_paragraph()
        _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
        run = para.add_run(authors)
        _set_font(run, italic=True, size=font_size)

    if institute:
        para = doc.add_paragraph()
        _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)
        run = para.add_run(institute)
        _set_font(run, italic=True, size=font_size - 1)

    # Separator line
    para = doc.add_paragraph()
    _set_para_format(para, space_before=4, space_after=8)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)

    fig_num = [0]
    section_cm = Cm(16)  # approximate page width for figures

    for block in blocks:
        if isinstance(block, HeadingBlock):
            size_map = {1: font_size+4, 2: font_size+2, 3: font_size+1, 4: font_size}
            bold_map = {1: True, 2: True, 3: False, 4: True}
            ital_map = {1: False, 2: False, 3: True, 4: False}
            before_map = {1: 20, 2: 14, 3: 10, 4: 8}
            after_map  = {1: 10, 2:  6, 3:  4, 4: 3}
            _add_text_para(block.text,
                            bold=bold_map.get(block.level, True),
                            italic=ital_map.get(block.level, False),
                            size=size_map.get(block.level, font_size),
                            align=WD_ALIGN_PARAGRAPH.CENTER if block.level == 1 else WD_ALIGN_PARAGRAPH.LEFT,
                            before=before_map.get(block.level, 10),
                            after=after_map.get(block.level, 4))

        elif isinstance(block, ParagraphBlock):
            _add_text_para(block.text, before=0, after=6)

        elif isinstance(block, DisplayMathBlock):
            png = render_math_png(block.expr, fontsize=font_size * 1.2, display=True)
            if png:
                para = doc.add_paragraph()
                _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8)
                run = para.add_run()
                run.add_picture(io.BytesIO(png), width=Inches(4.5))
            else:
                _add_text_para(_unicode_math(block.expr), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)

        elif isinstance(block, FigureBlock):
            if block.image_path and block.image_path.exists():
                fig_num[0] += 1
                para = doc.add_paragraph()
                _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=4)
                run = para.add_run()
                try:
                    run.add_picture(str(block.image_path), width=section_cm)
                except Exception:
                    pass
                cap_para = doc.add_paragraph()
                _set_para_format(cap_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=10)
                cap_run = cap_para.add_run(f"Figure {fig_num[0]}: {block.caption}")
                _set_font(cap_run, italic=True, size=font_size - 1)

        elif isinstance(block, TableBlock):
            if block.headers or block.rows:
                all_rows = ([block.headers] if block.headers else []) + block.rows
                if not all_rows: continue
                ncols = max(len(r) for r in all_rows)
                tbl = doc.add_table(rows=len(all_rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(all_rows):
                    for ci, cell_txt in enumerate(row):
                        if ci >= ncols: break
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = cell_txt
                        run = cell.paragraphs[0].runs
                        if run:
                            _set_font(run[0], bold=(ri == 0), size=font_size - 1)
                        else:
                            r = cell.paragraphs[0].add_run(cell_txt)
                            _set_font(r, bold=(ri == 0), size=font_size - 1)
                doc.add_paragraph()

        elif isinstance(block, ListBlock):
            for idx, item in enumerate(block.items):
                prefix = f"{idx+1}. " if block.ordered else "• "
                _add_text_para(f"{prefix}{item}", before=1, after=1)

        elif isinstance(block, RuleBlock):
            para = doc.add_paragraph()
            _set_para_format(para, space_before=6, space_after=6)

    doc.save(str(output_path))

# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "Usage: export_document.py <config_path>"}))
        sys.exit(1)

    config_path = Path(sys.argv[1])
    if not config_path.exists():
        print(json.dumps({"ok": False, "error": f"Config not found: {config_path}"}))
        sys.exit(1)

    try:
        config = json.loads(config_path.read_text())
    except Exception as e:
        print(json.dumps({"ok": False, "error": f"Bad config JSON: {e}"}))
        sys.exit(1)

    try:
        job_dir      = Path(config.get("job_dir", ""))
        output_path  = Path(config["output_path"])
        fmt          = config.get("format", "pdf").lower()
        content      = config.get("content", "")
        settings     = config.get("settings", {})
        fig_captions = config.get("figure_captions", {})

        output_path.parent.mkdir(parents=True, exist_ok=True)

        blocks = parse_content(content, job_dir if job_dir.exists() else None, fig_captions)

        if fmt == "docx":
            export_docx(blocks, output_path, settings)
        else:
            export_pdf(blocks, output_path, settings)

        print(json.dumps({"ok": True, "output_path": str(output_path), "format": fmt,
                          "blocks": len(blocks)}))
        sys.exit(0)
    except Exception as e:
        tb = traceback.format_exc()
        print(json.dumps({"ok": False, "error": str(e), "traceback": tb[-2000:]}))
        sys.exit(1)

if __name__ == "__main__":
    main()

"""
Document Parser — Docling-first, PyMuPDF fallback, pdfminer last resort
=========================================================================
Parses scientific PDFs into structured text with layout awareness.

Priority chain:
  1. Docling (IBM) — extracts tables, headings, captions, references
  2. PyMuPDF (fitz) — fast, layout-aware text + table detection
  3. pdfminer.six — basic text extraction fallback

Outputs a ParsedDocument with:
  - full_text: continuous text for regex/NIM extraction
  - sections: named sections (Abstract, Introduction, Methods, Results…)
  - tables: list of extracted tables as markdown strings
  - title / authors / abstract: auto-detected
"""

from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    full_text: str = ""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    abstract: str = ""
    sections: Dict[str, str] = field(default_factory=dict)
    tables: List[str] = field(default_factory=list)
    references: List[str] = field(default_factory=list)
    parser_used: str = "unknown"
    page_count: int = 0


# ── Section splitter ──────────────────────────────────────────────────────────

_SECTION_RE = re.compile(
    r"^(?:\d+\.?\s+)?("
    r"abstract|introduction|background|literature\s+review|"
    r"experimental|materials?\s+and\s+methods?|methodology|"
    r"results?|results?\s+and\s+discussion|discussion|"
    r"electrochemical\s+characterization|sensing\s+performance|"
    r"selectivity|interference|real\s+sample|food\s+sample|"
    r"conclusion[s]?|summary|acknowledgements?|references?"
    r")\s*$",
    re.IGNORECASE | re.MULTILINE,
)

_TITLE_RE = re.compile(r"^[A-Z][A-Za-z\s\-\(\),:/]{20,200}$", re.MULTILINE)
_ABSTRACT_SECTION_RE = re.compile(
    r"(?:abstract|summary)[:\s]*\n?(.*?)(?:\n{2,}|\bintroduction\b|\bkeywords?\b)",
    re.IGNORECASE | re.DOTALL,
)


def _split_sections(text: str) -> Dict[str, str]:
    """Split document text into named sections."""
    sections: Dict[str, str] = {}
    parts = _SECTION_RE.split(text)
    if len(parts) < 2:
        sections["body"] = text
        return sections

    # parts[0] = pre-abstract text, then alternating section-name / content
    sections["header"] = parts[0]
    for i in range(1, len(parts) - 1, 2):
        name = parts[i].lower().strip()
        content = parts[i + 1] if i + 1 < len(parts) else ""
        sections[name] = content.strip()

    return sections


def _infer_abstract(text: str) -> str:
    m = _ABSTRACT_SECTION_RE.search(text[:8000])
    if m:
        return m.group(1).strip()[:3000]
    # Fallback: first long paragraph
    paras = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 200]
    return paras[0][:1500] if paras else text[:800]


def _infer_title(text: str, filename: str = "") -> str:
    lines = [l.strip() for l in text[:3000].split("\n") if l.strip()]
    for line in lines[:15]:
        if 25 < len(line) < 250 and not line.startswith("http") and not re.match(r"^\d", line):
            return line
    return os.path.splitext(filename)[0] if filename else ""


def _infer_authors(text: str) -> List[str]:
    m = re.search(
        r"(?:authors?|by)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?(?:,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?){1,10})",
        text[:3000], re.IGNORECASE,
    )
    if m:
        return [a.strip() for a in m.group(1).split(",") if a.strip()][:8]
    return []


# ── 1. Docling parser ─────────────────────────────────────────────────────────

def _parse_with_docling(data: bytes, filename: str = "") -> Optional[ParsedDocument]:
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        opts = PdfPipelineOptions()
        opts.do_ocr = False
        opts.do_table_structure = True

        converter = DocumentConverter()
        buf = io.BytesIO(data)
        result = converter.convert(buf)
        doc = result.document

        # Full markdown export
        md_text = doc.export_to_markdown()
        full_text = re.sub(r"\[.*?\]\(.*?\)", "", md_text)  # strip links

        # Extract tables as markdown
        tables = []
        for table in doc.tables:
            try:
                tables.append(table.export_to_markdown())
            except Exception:
                pass

        doc_out = ParsedDocument(
            full_text=full_text,
            tables=tables,
            parser_used="docling",
            page_count=len(doc.pages) if hasattr(doc, "pages") else 0,
        )
        doc_out.title = _infer_title(full_text, filename)
        doc_out.abstract = _infer_abstract(full_text)
        doc_out.authors = _infer_authors(full_text)
        doc_out.sections = _split_sections(full_text)
        return doc_out

    except ImportError:
        return None
    except Exception as e:
        logger.warning("Docling parse failed (%s): %s", filename, e)
        return None


# ── 2. PyMuPDF parser ─────────────────────────────────────────────────────────

def _parse_with_pymupdf(data: bytes, filename: str = "") -> Optional[ParsedDocument]:
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(stream=data, filetype="pdf")
        pages_text = []
        tables_text = []

        for page in pdf:
            # Extract text with layout preservation
            text = page.get_text("text", sort=True)
            pages_text.append(text)

            # Extract tables
            try:
                tabs = page.find_tables()
                for tab in tabs.tables:
                    md = tab.to_markdown()
                    if md and len(md) > 30:
                        tables_text.append(md)
            except Exception:
                pass

        full_text = "\n".join(pages_text)
        pdf.close()

        doc_out = ParsedDocument(
            full_text=full_text,
            tables=tables_text[:20],
            parser_used="pymupdf",
            page_count=len(pages_text),
        )
        doc_out.title = _infer_title(full_text, filename)
        doc_out.abstract = _infer_abstract(full_text)
        doc_out.authors = _infer_authors(full_text)
        doc_out.sections = _split_sections(full_text)
        return doc_out

    except ImportError:
        return None
    except Exception as e:
        logger.warning("PyMuPDF parse failed (%s): %s", filename, e)
        return None


# ── 3. pdfminer fallback ──────────────────────────────────────────────────────

def _parse_with_pdfminer(data: bytes, filename: str = "") -> Optional[ParsedDocument]:
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams

        buf_in = io.BytesIO(data)
        buf_out = io.StringIO()
        extract_text_to_fp(buf_in, buf_out, laparams=LAParams(), output_type="text", codec=None)
        full_text = buf_out.getvalue()

        doc_out = ParsedDocument(
            full_text=full_text,
            parser_used="pdfminer",
        )
        doc_out.title = _infer_title(full_text, filename)
        doc_out.abstract = _infer_abstract(full_text)
        doc_out.authors = _infer_authors(full_text)
        doc_out.sections = _split_sections(full_text)
        return doc_out

    except ImportError:
        return None
    except Exception as e:
        logger.warning("pdfminer parse failed (%s): %s", filename, e)
        return None


# ── DOCX parser ───────────────────────────────────────────────────────────────

def _parse_docx(data: bytes, filename: str = "") -> Optional[ParsedDocument]:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_md = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows:
                tables_md.append("\n".join(rows))

        full_text = "\n".join(paragraphs)
        doc_out = ParsedDocument(
            full_text=full_text,
            tables=tables_md,
            parser_used="docx",
        )
        doc_out.title = _infer_title(full_text, filename)
        doc_out.abstract = _infer_abstract(full_text)
        doc_out.authors = _infer_authors(full_text)
        doc_out.sections = _split_sections(full_text)
        return doc_out

    except ImportError:
        return None
    except Exception as e:
        logger.warning("DOCX parse failed (%s): %s", filename, e)
        return None


# ── Public API ────────────────────────────────────────────────────────────────

def parse_document(
    data: bytes,
    mime_type: str,
    filename: str = "",
    prefer_docling: bool = True,
) -> ParsedDocument:
    """
    Parse a document (PDF, DOCX, or text) into a ParsedDocument.

    Parser priority for PDFs:
      1. Docling (if installed)
      2. PyMuPDF (fitz)
      3. pdfminer.six
    """
    is_pdf = "pdf" in mime_type or filename.lower().endswith(".pdf")
    is_docx = "wordprocessing" in mime_type or "docx" in mime_type or filename.lower().endswith(".docx")
    is_gdoc = "google-apps.document" in mime_type

    if is_docx or is_gdoc:
        result = _parse_docx(data, filename)
        if result and len(result.full_text.strip()) > 50:
            return result

    if is_pdf:
        # Try Docling first
        if prefer_docling:
            result = _parse_with_docling(data, filename)
            if result and len(result.full_text.strip()) > 50:
                logger.info("Parsed '%s' with Docling (%d chars, %d tables)",
                            filename, len(result.full_text), len(result.tables))
                return result

        # PyMuPDF
        result = _parse_with_pymupdf(data, filename)
        if result and len(result.full_text.strip()) > 50:
            logger.info("Parsed '%s' with PyMuPDF (%d chars, %d tables)",
                        filename, len(result.full_text), len(result.tables))
            return result

        # pdfminer fallback
        result = _parse_with_pdfminer(data, filename)
        if result and len(result.full_text.strip()) > 50:
            logger.info("Parsed '%s' with pdfminer (%d chars)", filename, len(result.full_text))
            return result

    # Plain text / generic
    try:
        text = data.decode("utf-8", errors="replace")
        doc_out = ParsedDocument(
            full_text=text,
            parser_used="plaintext",
        )
        doc_out.title = _infer_title(text, filename)
        doc_out.abstract = _infer_abstract(text)
        doc_out.sections = _split_sections(text)
        return doc_out
    except Exception:
        pass

    return ParsedDocument(full_text="", parser_used="failed")
